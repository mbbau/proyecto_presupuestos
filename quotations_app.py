import streamlit as st
import pandas as pd

st.set_page_config(
    page_title= "Quotations_project",
    layout= 'wide'
)

columna_logo = st.columns((0.7,0.3))
with columna_logo[1]:
    st.image('logo clinica.png')

st.header(body = "", divider='blue')

df = pd.DataFrame()

# Ahora genero las variables que serán ingresadas por el usuario al momento de generar el presupuesto.

columnas_para_datos_iniciales = st.columns((0.5,0.5))

with columnas_para_datos_iniciales[0]:
    presupuesto_estimado = st.text_input(
        "**Presupuesto Estimado NO.**",
        value = None
        )
    
with columnas_para_datos_iniciales[1]:
    fecha_de_presupuesto = st.date_input(
        "**Fecha de Presupuesto**", 
        value = "today"
        )

with columnas_para_datos_iniciales[0]:
    nombre_de_paciente = st.text_input(
        "**Nombre de Paciente**",
        value = None
        )
    
with columnas_para_datos_iniciales[1]:
    indentificacion_de_paciente = st.text_input(
        "**Identificación de Paciente**",
        value = None
        )

nombre_medico_futuro = st.markdown("**ACA IRA UN SELECT BOX PARA SELECCIONAR EL NOMBRE DEL MÉDICO**")
#medico = st.selectbox(
#    label = "Medico",
#    options = df.medico.unique(),
#    label_visibility = "hidden"
#)

## ACA DEBE IR CONVENIO.


columnas_procedimiento = st.columns((0.3,0.3,0.3))
with columnas_procedimiento[1]:
    st.markdown("**Presupuesto Estimado:**")
    st.markdown("**ACA IRA UN SELECT BOX PARA SELECCIONAR EL NOMBRE DEL PROCEDIMIENTO**")
#procedimiento = st.selectbox(
#    label = "Medico",
#    options = df2.procedimiento.unique(),
#    label_visibility = "hidden"
#)

st.markdown("""
        Estimado (a) cliente:  
        
         Para el Hospital Clínica Bíblica es un gusto saludarlo (a) y atender su requerimiento en cuanto a nuestros costos de servicios
         hospitalarios, en los cuales le ofrecemos una alta tecnología, profesionales de la salud debidamente calificados y una vasta
         experiencia de más de 93 años.
    
         Los costos de referencia para el procedimiento solicitado están basados en datos actualiados en cuanto a los últimos que se
         han realizado en nuestro hospital, y son los siguientes:

         """)



from docx import Document
from io import BytesIO

# Función para generar el documento de Word
def create_doc(medico, procedimiento, convenio):
    doc = Document()
    doc.add_heading('Presupuesto Médico', 0)
    doc.add_paragraph(f'Médico: {medico}')
    doc.add_paragraph(f'Procedimiento: {procedimiento}')
    doc.add_paragraph(f'Convenio: {convenio}')
    # Aquí podrías agregar más lógica para añadir contenido basado en los inputs
    return doc

# Función para convertir el documento a un stream de bytes
def convert_to_bytes(doc):
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()

# Crear la aplicación de Streamlit
st.title('Generador de Presupuestos Médicos')

# Recolectar inputs del usuario
medico = st.text_input('Nombre del Médico')
procedimiento = st.text_input('Nombre del Procedimiento')
convenio = st.text_input('Convenio (si aplica)')

if st.button('Generar Documento'):
    doc = create_doc(medico, procedimiento, convenio)
    doc_bytes = convert_to_bytes(doc)
    st.download_button(label="Descargar Presupuesto",
                       data=doc_bytes,
                       file_name="presupuesto.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")